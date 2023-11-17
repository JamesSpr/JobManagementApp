import shutil
import graphene
from graphql_jwt.decorators import login_required

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm


# import win32com.client as win32
# import pythoncom
from openpyxl import load_workbook

from datetime import datetime
import os
from ..models import Job, Estimate, EstimateHeader

import sys
sys.path.append("...")
from accounts.models import CustomUser


JOBS_PATH = r'C:\Users\Aurify Constructions\Aurify\Aurify - Maintenance\Jobs'

# class CreateBGISEstimate(graphene.Mutation):
#     class Arguments:
#         job_id = graphene.String()
#         selected_estimate = graphene.String()

#     success = graphene.Boolean()
#     message = graphene.String()

#     @classmethod
#     @login_required
#     def mutate(self, root, info, job_id, selected_estimate):
        
#         job = Job.objects.get(id=job_id)
#         estimate = Estimate.objects.get(job_id=job_id, name=selected_estimate)

#         # Check to see if file already exists
#         # This will have to be removed when we move to pdf only quotes
        
#         if os.path.exists(os.path.join(JOBS_PATH, str(job).strip(), "Estimates", str(estimate.name).strip(), "BGIS Estimate " + str(job).split(' ')[0] + ".xlsx")):
#             return self(success=False, message="Estimate Already Exists")

#         ## Create folder for the selected estimate
#         try:
#             os.mkdir(os.path.join(JOBS_PATH, str(job).strip(), "Estimates", str(estimate.name).strip()))
#         except FileExistsError:
#             pass

#         ## Save to new folder
#         shutil.copy(r"C:\Users\Aurify Constructions\Documents\JobManagementApp\project_management\api\templates\BGIS Estimate Template.xlsx", os.path.join(JOBS_PATH, str(job).strip(), "Estimates", str(estimate.name).strip(), "BGIS Estimate " + str(job).split(' ')[0] + ".xlsx"))

#         xl_file = os.path.join(JOBS_PATH, str(job).strip(), "Estimates", str(estimate.name).strip(), "BGIS Estimate " + str(job).split(' ')[0] + ".xlsx")
#         workbook = load_workbook(filename=xl_file)
#         sheet = workbook.active
#         sheet["C5"] = estimate.name
#         workbook.save(filename=xl_file) 

#         # xlApp = win32.DispatchEx("Excel.Application", pythoncom.CoInitialize())
#         # xlApp.Visible = False
#         # wb = xlApp.Workbooks.Open(os.path.join(JOBS_PATH, str(job).strip(), "Estimates", str(estimate.name).strip(), "BGIS Estimate " + str(job).split(' ')[0] + ".xlsx"))
#         # ws = wb.Sheets("Cost Breakdown")
#         # wb.Close(True)
#         # xlApp.Quit()
#         # del xlApp

#         return self(success=True, message="Estimate Created Successfully")


class CreateCompletionDocuments(graphene.Mutation):
    class Arguments:
        job_id = graphene.String()
        
    success = graphene.Boolean()
    message = graphene.String()

    @classmethod
    @login_required
    def mutate(self, root, info, job_id):

        job = Job.objects.get(id=job_id)
        templates_path = r'C:\Users\Aurify Constructions\Documents\JobManagementApp\project_management\api\templates'

        if job.po == "" or job.po == None:
            return self(success=False, message="Please ensure job has a PO Number")

        if job.scope == "" or job.scope == None:
            return self(success=False, message="Please ensure job has a Scope of Works")

        if job.site_manager == "" or job.site_manager == None:
            return self(success=False, message="Please ensure a site manager is selected")

        if not os.path.exists(os.path.join(JOBS_PATH, str(job).strip())):
            return self(success=False, message="Job folder does not exist")
        
        if not os.path.exists(os.path.join(JOBS_PATH, str(job).strip(), "Documentation")):
            return self(success=False, message="File System Folders are not correct. Please check Documentation Folder Exists")

        # word = win32.DispatchEx("Word.Application", pythoncom.CoInitialize())
        # word.Visible = False

        table_items = {
            '[Job]': str(job),
            '[JobNumber]': job.po,
            '[JobTitle]': job.title,
            '[SiteManager]': job.site_manager.first_name.title().strip() + " " + job.site_manager.last_name.capitalize().strip(),
            '[Date]': datetime.now().strftime("%d/%m/%Y"), 
            '[Address]': job.location.getFullAddress(),
            '[Scope]': job.scope,
        }

        paragraph_items = {
            '[JobNumber]': job.po,
            '[Scope]': job.scope,
            '[Date]': datetime.now().strftime("%d/%m/%Y"), 
        }

        document = None

        swms_filename = os.path.join(JOBS_PATH, str(job).strip(), "Documentation", job.po + " - SWMS.docx")
        pra_filename = os.path.join(JOBS_PATH, str(job).strip(), "Documentation", job.po + " - PRA.docx")
        sdkt_filename = os.path.join(JOBS_PATH, str(job).strip(), "Documentation", job.po + " - SDKT.docx")

        try:
            if not os.path.exists(os.path.join(JOBS_PATH, str(job).strip(), "Documentation", job.po + " - SWMS.docx")):
                if not job.site_manager == None or not job.site_manager == "":
                    # Open SWMS
                    document = Document(os.path.join(templates_path, "SWMS.docx"))
                    # document = word.Documents.Add(Template=templates_path + "/SWMS.dotx", NewTemplate=False, DocumentType=0)
                    
                    for item in table_items:
                        for table in document.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        if item in paragraph.text:
                                            paragraph.font = document.styles['Normal']
                                            paragraph.text = paragraph.text.replace(item, table_items[item])
                                 

                    # Save and close word document
                    document.save(swms_filename)

                else:
                    return self(success=False, message="Site Manager Required for SWMS Creation")
                
            if not os.path.exists(os.path.join(JOBS_PATH, str(job).strip(), "Documentation", job.po + " - PRA.docx")):
                if not job.site_manager == None or not job.site_manager == "":
                    # Open PRA
                    document = Document(os.path.join(templates_path, "PRA.docx"))
                    
                    for item in table_items:
                        for table in document.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for paragraph in cell.paragraphs:
                                        if item in paragraph.text:
                                            paragraph.font = document.styles['Normal']
                                            paragraph.text = paragraph.text.replace(item, table_items[item])



                    # Save and close word document
                    document.save(pra_filename)

                else:
                    return self(success=False, message="Site Manager Required for Pre-Start Risk Assessment Creation")
            
            if not os.path.exists(os.path.join(JOBS_PATH, str(job).strip(), "Documentation", job.po + " - SDKT.docx")):
                if not job.completion_date == "" or not job.completion_date == None:
                    # Open Service Docket
                    document = Document(os.path.join(templates_path, "SDKT.docx"))

                    for item in paragraph_items:
                        for paragraph in document.paragraphs:
                            if item in paragraph.text:
                                paragraph.text = paragraph.text.replace(item, paragraph_items[item])

                    # Save and close word document
                    document.save(sdkt_filename)
                    
                else:
                    return self(success=False, message="Completion Date required for Service Docket.")
                
        except BaseException as err:
            print(err)
            print(swms_filename, "\n", pra_filename, "\n", sdkt_filename)
            return self(success=False, message="An Error has occured, please contact admin.")

        return self(success=True, message="Completion Documents Saved to Folder. Please fill out the SWMS and PRA")