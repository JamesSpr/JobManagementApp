import shutil
import graphene

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Cm, Pt

from datetime import date
import os
from ..models import Job, Estimate, EstimateHeader

import sys
sys.path.append("...")
from accounts.models import CustomUser

QUOTE_PATH = "api/templates/QUOTE.docx"
PO_PATH = "api/templates/PURCHASE ORDER.docx"
JOBS_PATH = r'C:\Users\Aurify Constructions\Aurify\Aurify - Maintenance\Jobs'

class CreateQuote(graphene.Mutation):
    class Arguments:
        job_id = graphene.String()
        selected_estimate = graphene.String()
        userId = graphene.String()

    success = graphene.Boolean()
    message = graphene.String()

    @classmethod
    def mutate(cls, root, info, job_id, selected_estimate, userId):
        
        job = Job.objects.get(id=job_id)
        estimate = Estimate.objects.get(job_id=job_id, id=selected_estimate)
        estimate_headers = EstimateHeader.objects.filter(estimate_id=estimate)
        user = CustomUser.objects.get(id=userId)
        # print("Creating Quote for", job)
        # print("Selected Estimate", estimate.name)
    
        # Check to see if file already exists# Save to dropbox
        # This will have to be removed when we move to pdf only quotes
        if os.path.exists(os.path.join(JOBS_PATH, str(job).strip(), "Estimates", str(estimate.name).strip(), 'Aurify Quote ' + str(estimate.name).strip() + '.docx')):
            return CreateQuote(success=False, message="Quote Already Exists")

        # Create File
        document = Document(os.path.join(os.getcwd(), QUOTE_PATH))

        # Get data from job
        quote_date = date.today().strftime("%d/%m/%Y").strip()
        estimate_name = estimate.name
        client_name = job.client.name.strip()
        address = job.location.address.strip()
        locality = job.location.locality.strip()
        state = job.location.state.strip()
        postcode = job.location.postcode.strip()
        job_name = str(job).strip()
        location = job.location.name.strip()
        scope = estimate.scope.strip()

        contact_name = job.requester.first_name.strip() + " " + job.requester.last_name.strip()
        contact_email = job.requester.email.strip()
        contact_phone = job.requester.phone.strip()

        # Fill File
        mailMergeTableItems = {
            '[QuoteDate]': quote_date, 
            '[EstimateDescription]': estimate_name, 
            '[ClientName]': client_name, 
            '[Address]': address,
            '[Locality]': locality, 
            '[State]': state, 
            '[Postcode]': postcode, 
            '[ContactName]': contact_name, 
            '[ContactEmail]': contact_email.capitalize().strip(), 
            '[ContactPhone]': contact_phone,
            '[JobName]': job_name,
            '[QuoterPhone]': estimate.quote_by.phone.strip(),
            '[QuoterEmail]': estimate.quote_by.email.capitalize().strip(),
        }

        mailMergeParagraphItems = {
            '[ClientName]' : client_name, 
            '[Location]': location,
            '[Scope]': scope,
            '[QuoterSignature]': rf'C:\Users\Aurify Constructions\Documents\JobManagementApp\project_management\api\templates\temp\user_{estimate.quote_by.first_name.lower()}-signature.png',
            '[CompanyLogo]' : r'C:\Users\Aurify Constructions\Documents\JobManagementApp\project_management\api\templates\temp\aurify_logo-quote.png',
            '[Quoter]': estimate.quote_by.first_name.title().strip() + " " + estimate.quote_by.last_name.capitalize().strip(),
            '[QuoterRole]': estimate.quote_by.position.title().strip(),
        }

        for item in mailMergeTableItems:
            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if item in paragraph.text:
                                paragraph.font = document.styles['Normal']
                                paragraph.text = paragraph.text.replace(item, mailMergeTableItems[item])
                                # print(paragraph.text)

        for item in mailMergeParagraphItems:
            for paragraph in document.paragraphs:
                if item in paragraph.text:
                    if item == '[QuoterSignature]':
                        paragraph.text = paragraph.text.replace(item, '')
                        paragraph.left_indent = 0.5
                        sig_width = 2.25 if estimate.quote_by.first_name.lower() == "mark" else 4.75
                        paragraph.add_run().add_picture(mailMergeParagraphItems[item],  width=Cm(sig_width))
                    elif item == '[CompanyLogo]':
                        paragraph.text = paragraph.text.replace(item, '')
                        paragraph.left_indent = 0.5
                        paragraph.add_run().add_picture(mailMergeParagraphItems[item], width=Cm(5.4))
                    else:
                        paragraph.text = paragraph.text.replace(item, mailMergeParagraphItems[item])
                        paragraph.left_indent = 0.5

        # Create estimate table
        for table in document.tables:
            # Find correct table
            if table.rows[1].cells[0].paragraphs[0].text == "[TradesWithSummary]":
                table.rows[1].cells[0].paragraphs[0].text.replace("[TradesWithSummary]", "")

                # Add new row for each header item
                for index, header in enumerate(estimate_headers):
                    # print(index+1, header)
                    table.rows[index+1].cells[0].text = f"{index+1:02d}"
                    table.rows[index+1].cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    table.rows[index+1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.rows[index+1].cells[1].text = header.description
                    table.rows[index+1].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    table.rows[index+1].cells[2].text = str("{:,.2f}".format(header.gross))
                    table.rows[index+1].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    table.rows[index+1].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    table.rows[index+1].height = Cm(0.8)
                    table.add_row()

                # Add total rows
                table.add_row()
                table.add_row()
                table.rows[len(estimate_headers) + 1].height = Cm(0.8)
                table.rows[len(estimate_headers) + 1].cells[0].merge(table.rows[len(estimate_headers) + 1].cells[1])
                table.rows[len(estimate_headers) + 1].cells[1].paragraphs[0].add_run(text = "Total (excl. GST)").bold = True
                table.rows[len(estimate_headers) + 1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                table.rows[len(estimate_headers) + 1].cells[2].paragraphs[0].add_run(text = "$" + str("{:,.2f}".format(estimate.price))).bold = True
                table.rows[len(estimate_headers) + 1].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.rows[len(estimate_headers) + 1].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                table.rows[len(estimate_headers) + 1].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                table.rows[len(estimate_headers) + 2].height = Cm(0.8)
                table.rows[len(estimate_headers) + 2].cells[0].merge(table.rows[len(estimate_headers) + 2].cells[1])
                table.rows[len(estimate_headers) + 2].cells[1].paragraphs[0].add_run(text = "GST").bold = True 
                table.rows[len(estimate_headers) + 2].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                table.rows[len(estimate_headers) + 2].cells[2].paragraphs[0].add_run(text =  "$" + str("{:,.2f}".format(round(float(estimate.price) * 0.1, 2)))).bold = True 
                table.rows[len(estimate_headers) + 2].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.rows[len(estimate_headers) + 2].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                table.rows[len(estimate_headers) + 2].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                table.rows[len(estimate_headers) + 3].height = Cm(0.8)
                table.rows[len(estimate_headers) + 3].cells[0].merge(table.rows[len(estimate_headers) + 3].cells[1])
                table.rows[len(estimate_headers) + 3].cells[1].paragraphs[0].add_run(text = "Total (incl. GST)").bold = True 
                table.rows[len(estimate_headers) + 3].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                table.rows[len(estimate_headers) + 3].cells[2].paragraphs[0].add_run(text = "$" + str("{:,.2f}".format(round(float(estimate.price) * 1.1, 2)))).bold = True
                table.rows[len(estimate_headers) + 3].cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                table.rows[len(estimate_headers) + 3].cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                table.rows[len(estimate_headers) + 3].cells[2].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            

        ## Create folder for the selected estimate
        try:
            os.mkdir(os.path.join(JOBS_PATH, str(job).strip(), "Estimates", str(estimate.name).strip()))
        except FileExistsError:
            pass
        ## Save to new folder
        document.save(os.path.join(JOBS_PATH, str(job).strip(), "Estimates", str(estimate.name).strip(), 'Aurify Quote ' + str(estimate.name).strip() + '.docx'))

        ## Save to local for backup
        document.save(os.path.join(os.getcwd(),'api/generated_quotes/Aurify Quote ' + str(estimate.name).strip() + '.docx'))

        return CreateQuote(success=True, message="Quote Created Successfully")

import win32com.client as win32
import pythoncom

class CreateBGISEstimate(graphene.Mutation):
    class Arguments:
        job_id = graphene.String()
        selected_estimate = graphene.String()

    success = graphene.Boolean()
    message = graphene.String()

    @classmethod
    def mutate(cls, root, info, job_id, selected_estimate):
        
        job = Job.objects.get(id=job_id)
        estimate = Estimate.objects.get(job_id=job_id, id=selected_estimate)

        # Check to see if file already exists
        # This will have to be removed when we move to pdf only quotes
        if os.path.exists(os.path.join(JOBS_PATH, str(job).strip(), "Estimates", str(estimate.name).strip(), "BGIS Estimate " + str(estimate.name).strip() + ".xlsx")):
            return cls(success=False, message="Estimate Already Exists")

        ## Create folder for the selected estimate
        try:
            os.mkdir(os.path.join(JOBS_PATH, str(job).strip(), "Estimates", str(estimate.name).strip()))
        except FileExistsError:
            pass

        ## Save to new folder
        # shutil.copy(r"C:\Users\Aurify Constructions\Documents\JobManagementApp\project_management\api\templates\BGIS Estimate Template.xlsx", os.path.join(JOBS_PATH, str(job).strip(), "Estimates", str(estimate.name).strip(), "BGIS Estimate " + str(estimate.name).strip() + ".xlsx"))

        xlApp = win32.DispatchEx("Excel.Application", pythoncom.CoInitialize())
        xlApp.Visible = True
        wb = xlApp.Workbooks.Open(r"C:\Users\Aurify Constructions\Documents\JobManagementApp\project_management\api\templates\BGIS Estimate Template.xlsx")
        ws = wb.Sheets("Cost Breakdown")
        ws.Range("C5").Value = estimate.name 

        wb.SaveAs(os.path.join(JOBS_PATH, str(job).strip(), "Estimates", str(estimate.name).strip(), "BGIS Estimate " + str(estimate.name).strip() + ".xlsx"), 51)
        wb.Close()
        xlApp.Quit()
        del xlApp

        return cls(success=True, message="Estimate Created Successfully")